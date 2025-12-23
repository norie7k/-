
from __future__ import annotations
import json, time, typing as T
import pandas as pd
import requests
from pathlib import Path
from typing import List, Dict, Any
import re, json, unicodedata
import json

# --- openpyxl 样式/工具 ---
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side, NamedStyle

# --- docx 样式/工具 ---
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


################模型调用，出结果###################

def load_system_prompt(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"Prompt file not found: {path}")
    return path.read_text(encoding="utf-8")

def build_user_prompt_filter(json_lines: T.List[str]) -> str:
    # 模型#1：筛掉非游戏相关，只输出相关 JSON 行（原样）
    return (
        "以下是若干玩家/客服/研发的发言记录，请根据系统提示中规则，"
        "判断哪些是【与游戏内容相关】的发言，保留这些 JSON 行，不相关的忽略。"
        "请仅输出【相关发言的原始 JSON 行】，严格保持格式不变。\n\n"
        "【输入】：\n" + "\n".join(json_lines)
    )

def build_user_prompt_discuss(jsonl_block: str) -> str:
    return (
        "- 请按以下 JSON 格式输出，并根据输入自动决定事件数量；"
        "不要生成任何额外说明文字，也不要固定输出两个事件。\n\n"
        
        "【输出格式示例】（仅示例，不要重复或固定事件数量）：\n"
        "[\n"
        "  {\n"
        "    \"事件名称\": \"示例事件\",\n"
        "    \"核心讨论摘要\": \"事件的核心讨论摘要总结\",\n"
        "    \"讨论热度（量化）\": \"发言玩家总数：X 位，发言总量：Y 条\"\n"
        "  }\n"
        "]\n\n"

        "【要求】\n"
        "1. 输出为 JSON 数组\n"
        "2. 数组中每个对象代表一个独立事件\n"
        "3. 事件数量根据输入数据自动生成，不允许固定两个或重复模板\n"
        "4. 不输出多余说明文字\n\n"

        "【输入】：\n" + jsonl_block
    )


def build_user_prompt_emotion(jsonl_block: str) -> str:
    return (
        "请按以下 JSON 格式输出，并根据输入自动决定事件数量；"
        "不要生成任何额外说明文字，也不要固定事件数量。\n\n"
        "【输出格式示例】（仅示例，不要重复或固定事件数量）：\n"
        "[\n"
        "  {\n"
        "    \"情绪事件\": \"示例事件\",\n"
        "    \"简要情绪总结\": \"精炼概括总结玩家的情绪表现特点\",\n"
        "    \"情绪波动类型\": \"负向爆发 / 正向爆发 / 落差感 / 情绪跳变\",\n"
        "    \"情绪趋势\": \"X → Y\",\n"
        "    \"情绪热度（量化）\": \"发言玩家总数：X 位，发言总量：Y 条\",\n"
        "    \"代表性情绪发言示例\": [\"示例句1\", \"示例句2\"]\n"
        "  }\n"
        "]\n\n"
        "【输入】：\n" + jsonl_block
    )

def build_ingame_special(jsonl_block: str) -> str:
    return (
        "请按以下 JSON 格式输出，并根据输入自动决定事件数量；"
        "不要生成任何额外说明文字，也不要固定事件数量。\n\n"
        "【输出格式示例】（仅示例，不要重复或固定事件数量）：\n"
        "[\n"
        "  {\n"
        "    \"事件名称\": \"示例事件\",\n"
        "    \"事件摘要\": \"精炼概括总结归纳玩家在该事件中的表达焦点与讨论内容\",\n"
        "    \"玩家共识\": \"提取集体对事件的主流看法或共同认识\",\n"
        "    \"事件影响\": \"明确识别该事件对玩家带来的实际行为层面的影响\",\n"
        "    \"讨论热度（量化）\": \"发言玩家总数：X 位，发言总量：Y 条\",\n"
        "  }\n"
        "]\n\n"
        "【输入】：\n" + jsonl_block
    )

def build_outgame_special(jsonl_block: str) -> str:
    return (
        "请按以下 JSON 格式输出，并根据输入自动决定事件数量；"
        "不要生成任何额外说明文字，也不要固定事件数量。\n\n"
        "【输出格式示例】（仅示例，不要重复或固定事件数量）：\n"
        "[\n"
        "  {\n"
        "    \"事件类型\": \"属于什么类型公共经历事件\",\n"
        "    \"事件名称\": \"事件名称\",\n"
        "    \"事件摘要\": \"精炼概括总结基于该特殊事件下玩家整体的共识看法\",\n"
        "    \"玩家共识\": \"明确识别该事件对玩家带来的实际行为层面的影响\",\n"
        "    \"事件影响\": \该事件对游戏相关讨论、社群节奏、氛围带来了什么影响？\",\n"
         "   \"讨论热度（量化）\": \"发言玩家总数：X 位，发言总量：Y 条\",\n"
        "  }\n"
        "]\n\n"
        "【输入】：\n" + jsonl_block
    )


def call_ark_chat_completions(
    api_url: str,
    api_key: str,
    model: str,
    system_prompt: str,
    user_prompt: str,
    temperature: float = 0.3,
    max_tokens: int = 32700,
    timeout: int = 600,
    retries: int = 2,
) -> str:
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    last_err = None
    for attempt in range(retries + 1):
        try:
            resp = requests.post(api_url, headers=headers, json=payload, timeout=timeout)
            if resp.status_code != 200:
                raise RuntimeError(f"HTTP {resp.status_code}: {resp.text}")
            data = resp.json()
            return data["choices"][0]["message"]["content"]
        except Exception as e:
            last_err = e
            time.sleep(1.2 * (attempt + 1))
    raise RuntimeError(f"Ark API 调用失败: {last_err}")


from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Pt
from pathlib import Path
import json

########## 创建word文档 ##########
def create_word_report_grouped(path: Path):
    doc = Document()

    # 主标题
    title = doc.add_heading("玩家发言分析报告", level=0)
    r = title.runs[0]
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(22)
    r.bold = True

    doc.add_paragraph("")

    # 一级章节：热度讨论
    doc.add_heading("一、热度讨论", level=1)

    # 一级章节：情绪高点分析
    doc.add_heading("二、情绪高点分析", level=1)

    doc.save(path)


########## 工具：在某段落后插入新段落 ##########
def insert_paragraph_after(paragraph, text="", style=None):
    """
    在给定 paragraph 后面插入一个新的段落，返回这个新段落
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    parent = paragraph._parent

    # 先占位一个段落，再把底层 xml 换成新节点
    new_para = parent.add_paragraph()
    new_para._p = new_p

    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)

    return new_para


########## 热度事件写入（聚合到“一、热度讨论”下） ##########
def append_hot_grouped(path: Path, content: str):
    doc = Document(path)

    # 找到标题“ 一、热度讨论 ”
    target = None
    for p in doc.paragraphs:
        if "热度讨论" in p.text:   # 用关键词匹配更稳，不纠结前面的“一、”
            target = p
            break

    if target is None:
        raise ValueError("找不到标题：一、热度讨论")

    # 解析 JSON
    try:
        events = json.loads(content)
    except Exception:
        start = content.find("[")
        end = content.rfind("]") + 1
        events = json.loads(content[start:end])

    current = target  # 当前插入锚点

    for idx, ev in enumerate(events, start=1):
        # ===== 事件标题 =====
        current = insert_paragraph_after(current)
        r = current.add_run(f"{idx}. {ev.get('事件名称', '')}")
        r.bold = True
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt(12)
        current.paragraph_format.line_spacing = 1.5

        # ===== 核心讨论摘要 =====
        if "核心讨论摘要" in ev:
            p = insert_paragraph_after(current, style="List Bullet")
            r1 = p.add_run(f"核心讨论摘要：{ev['核心讨论摘要']}")
            r1.font.name = "黑体"
            r1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            r1.font.size = Pt(12)
            p.paragraph_format.line_spacing = 1.5
            current = p

        # ===== 热度 =====
        if "讨论热度（量化）" in ev:
            p = insert_paragraph_after(current, style="List Bullet")
            r2 = p.add_run(f"讨论热度：{ev['讨论热度（量化）']}")
            r2.font.name = "黑体"
            r2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            r2.font.size = Pt(12)
            p.paragraph_format.line_spacing = 1.5
            current = p

        # 空一行
        current = insert_paragraph_after(current, "")

    doc.save(path)


########## 情绪高点事件写入（聚合到“二、情绪高点分析”下） ##########
def append_emotion_grouped(path: Path, content: str):
    doc = Document(path)

    target = None
    for p in doc.paragraphs:
        if "情绪高点分析" in p.text:
            target = p
            break

    if target is None:
        raise ValueError("找不到标题：二、情绪高点分析")

    try:
        events = json.loads(content)
    except Exception:
        start = content.find("[")
        end = content.rfind("]") + 1
        events = json.loads(content[start:end])

    current = target

    for idx, ev in enumerate(events, start=1):
        # ===== 情绪事件标题 =====
        current = insert_paragraph_after(current)
        r = current.add_run(f"{idx}. {ev.get('情绪事件', '')}")
        r.bold = True
        r.font.name = "黑体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r.font.size = Pt(12)
        current.paragraph_format.line_spacing = 1.5

        # ===== 各字段 =====
        for key in ["简要情绪总结", "情绪波动类型", "情绪趋势", "情绪热度（量化）"]:
            if key in ev:
                p = insert_paragraph_after(current, style="List Bullet")
                r1 = p.add_run(f"{key}：{ev[key]}")
                r1.font.name = "黑体"
                r1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                r1.font.size = Pt(12)
                p.paragraph_format.line_spacing = 1.5
                current = p

        current = insert_paragraph_after(current, "")

    doc.save(path)


